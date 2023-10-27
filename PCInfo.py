import os
import socket
import subprocess
import sys

import pandas as pd
from docx.enum.section import WD_SECTION, WD_ORIENT
from getmac import get_mac_address
import wmi
import tkinter as tk
import re
from pandastable import Table
import psutil
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor
from tkinter import messagebox, Tk, filedialog
from PIL import Image, ImageTk
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils.dataframe import dataframe_to_rows
from screeninfo import get_monitors


def save_to_excel(df):
    # Crea una nueva instancia de Workbook y selecciona la hoja activa
    wb = Workbook()
    ws = wb.active

    # Ajusta el inicio de la tabla en la celda B8
    for r, row in enumerate(dataframe_to_rows(df, index=False, header=True), 8):
        for c, value in enumerate(row, 2):
            ws.cell(row=r, column=c, value=value)

    # Ajusta el tamaño de las celdas al contenido
    for column in ws.columns:
        max_length = 0
        column = [cell for cell in column]
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(cell.value)
            except:
                pass
        adjusted_width = (max_length + 2)
        ws.column_dimensions[column[0].column_letter].width = adjusted_width

    # Añade color gris a los encabezados
    gray_fill = PatternFill(start_color="D3D3D3",
                            end_color="D3D3D3",
                            fill_type="solid")
    gray_fill = PatternFill(start_color="D3D3D3",
                            end_color="D3D3D3",
                            fill_type="solid")
    for cells in ws["B8:R8"]:
        for cell in cells:
            cell.fill = gray_fill

    # Define los bordes
    thin_border = Border(left=Side(style='thin'),
                         right=Side(style='thin'),
                         top=Side(style='thin'),
                         bottom=Side(style='thin'))

    for row in ws.iter_rows(min_row=8, min_col=2, max_row=df.shape[0] + 8, max_col=df.shape[1] + 2):
        for cell in row:
            cell.border = thin_border
    # Centra el contenido de todas las celdas de la tabla
    for row in ws.iter_rows(min_row=8, min_col=2, max_row=df.shape[0] + 8, max_col=df.shape[1] + 2):
        for cell in row:
            cell.alignment = Alignment(horizontal='center', vertical='center')

    # Combina las celdas de B4 hasta R4 y añade el texto en negritas
    ws.merge_cells('B4:R4')
    ws['B4'].font = Font(bold=True)
    ws['B4'] = "LEVANTAMIENTO DE INFORMACIÓN DE EQUIPOS DE COMPUTO"
    ws['B4'].alignment = Alignment(horizontal='center', vertical='center')

    # Guarda el workbook
    wb.save('DIRECCION DE TECNOLOGIAS.xlsx')


def get_system_info():
    computer = wmi.WMI().Win32_ComputerSystem()[0]
    bios = wmi.WMI().Win32_BIOS()[0]
    processor = wmi.WMI().Win32_Processor()[0]
    ram = round(psutil.virtual_memory().total / (1024.0 ** 3))
    disk = round(psutil.disk_usage('/').total / (1024.0 ** 3))

    info = {'Tipo': 'CPU', 'Marca': computer.Manufacturer, 'Modelo': computer.Model, 'Serie': bios.SerialNumber,
            'Caracteristicas': f"{processor.NumberOfCores} CORE {processor.Name}, {processor.MaxClockSpeed} MHZ {ram} GB RAM DD {disk}GB",
            'Inf. en GB': get_disk_info(), 'MAC': get_mac_address()}
    hostname = socket.gethostname()
    info['IP'] = socket.gethostbyname(hostname)
    return info


def get_disk_info():
    disk_info = []
    for disk in psutil.disk_partitions():
        if 'fixed' in disk.opts and disk.device == 'C:\\':  # Solo el disco C:
            usage = psutil.disk_usage(disk.mountpoint)
            total_gb = round(usage.total / (2 ** 30), 1)  # Convierte bytes a GB
            used_gb = round(usage.used / (2 ** 30), 1)  # Convierte bytes a GB
            disk_info.append(f"C: {total_gb} GB ({used_gb} GB en uso)")

    return ', '.join(disk_info)


#def get_monitor_info():
 #   c = wmi.WMI(namespace="root\wmi")
  #  monitors = c.WmiMonitorID()
    #info = {}
   # for monitor in monitors:
     #   info['Tipo'] = 'Monitor'
      #  info['Marca'] = "".join(chr(i) for i in monitor.ManufacturerName) if monitor.ManufacturerName else 'Desconocido'
       # info['Modelo'] = "".join(
        #    chr(i) for i in monitor.UserFriendlyName) if monitor.UserFriendlyName else 'Desconocido'
        #info['Serie'] = "".join(chr(i) for i in monitor.SerialNumberID) if monitor.SerialNumberID else 'Desconocido'
    #return info
def get_monitor_info():
    c = wmi.WMI(namespace="root\wmi")
    monitors = c.WmiMonitorID()
    monitor_count = len(get_monitors())  # Contar el número de monitores
    info = []
    for monitor in monitors:
        monitor_info = {}
        monitor_info['Tipo'] = 'Monitor'
        monitor_info['Marca'] = "".join(chr(i) for i in monitor.ManufacturerName) if monitor.ManufacturerName else 'Desconocido'
        monitor_info['Modelo'] = "".join(chr(i) for i in monitor.UserFriendlyName) if monitor.UserFriendlyName else 'Desconocido'
        monitor_info['Serie'] = "".join(chr(i) for i in monitor.SerialNumberID) if monitor.SerialNumberID else 'Desconocido'
        if monitor_info['Marca'] == "HWP":
            monitor_info['Marca'] = "HP"

        info.append(monitor_info)
    return info, monitor_count

def clean_text(text):
    return re.sub(r'[\x00-\x1f\x7f-\x9f]', '', text)


def show_df(df):
    root = tk.Tk()
    frame = tk.Frame(root)
    frame.pack(fill='both', expand=True)
    pt = Table(frame, dataframe=df)
    pt.show()

    root.mainloop()


def reorder_columns(df):
    df = df[
        ['No.', 'Departamento', 'Nombre de Usuario', 'Correo Corporativo', 'Cargo', 'Tipo', 'Marca', 'Modelo', 'Serie',
         'Inventario', 'Caracteristicas', 'Inf. en GB', 'MAC', 'IP']]
    return df


def read_existing_document():
    document = Document('DIRECCION DE TECNOLOGIAS.docx')
    table = document.tables[-1]

    data = []
    for i, row in enumerate(table.rows):
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text)
        # Considera la primera fila como los nombres de las columnas
        if i == 0:
            columns = [name.strip() for name in row_data]  # Elimina los espacios al final
        else:
            data.append(row_data)

    df_existing = pd.DataFrame(data, columns=columns)

    return df_existing


def auto_increment(df):
    df['No.'] = pd.to_numeric(df['No.'], errors='coerce')
    if 'No.' in df.columns and df['No.'].max() >= 1:
        start = int(df['No.'].max()) + 1
        stop = start + len(df)
        df['No.'] = range(start, stop)
    else:
        df['No.'] = range(1, len(df) + 1)

    return df


def configure_margins(doc):
    section = doc.sections[0]
    section.top_margin = Pt(10)
    section.bottom_margin = Pt(10)
    section.left_margin = Pt(10)
    section.right_margin = Pt(10)
    # Cambiar la orientación a paisaje
    section.orientation = WD_ORIENT.LANDSCAPE

    # Ajustar el tamaño de la página a la orientación
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height


def add_title(doc):
    title = doc.add_heading(level=1)
    run = title.add_run('FORMATO PARA RECOPILACIÓN DE INFORMACIÓN DE EQUIPOS')
    run.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table_to_doc(df, doc):
    t = doc.add_table(df.shape[0] + 1, df.shape[1])

    for j in range(df.shape[-1]):
        t.cell(0, j).text = df.columns[j]

    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i + 1, j).text = clean_text(str(df.values[i, j]))

    return t


def stylize_table(t):
    t.style = 'Table Grid'
    for row in t.rows:
        for cell in row.cells:
            cell.width = Pt(0)


def stylize_table_header(t):
    for cell in t.row_cells(0):
        cell.text = cell.text.upper()
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(255, 255, 255)  # Color blanco
        cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w'))))  # Fondo negro


class SplashScreen(object):
    def __init__(self, image_file, duration=3):
        self.root = tk.Tk()
        self.root.overrideredirect(True)

        # Centra la ventana de splash
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # Obtiene la ruta al directorio del ejecutable
        if getattr(sys, 'frozen', False):
            application_path = sys._MEIPASS
        else:
            application_path = os.path.dirname(os.path.abspath(__file__))

        # Abre la imagen
        image_path = os.path.join(application_path, image_file)
        image = Image.open(image_path)  # Aquí es donde se abre la imagen
        width, height = image.size

        # Convierte la imagen en un objeto PhotoImage que tkinter puede usar
        self.image = ImageTk.PhotoImage(image)

        x = (screen_width // 2) - (width // 2)
        y = (screen_height // 2) - (height // 2)

        self.root.geometry(f"{width}x{height}+{x}+{y}")

        # Crea un lienzo y dibuja la imagen
        canvas = tk.Canvas(self.root, width=width, height=height)
        canvas.create_image(0, 0, image=self.image, anchor='nw')
        canvas.pack()

        # Muestra la ventana de splash durante un tiempo determinado y luego la destruye
        self.root.after(duration * 1000, self.root.destroy)

    def show(self):
        # Inicia el bucle principal de tkinter
        self.root.mainloop()


def clean_illegal_chars(df):
    for col in df.columns:
        if df[col].dtype == 'object':
            df[col] = df[col].apply(lambda x: ''.join([" " if ord(i) < 32 or ord(i) > 126 else i for i in str(x)]))
    return df


def main():
    splash = SplashScreen("logo_ancho.jpg")
    splash.show()
    departamento = input("Ingrese el departamento: ")
    nombre_usuario = input("Ingrese el nombre del usuario: ")
    cargo = input("Ingrese el cargo del usuario: ")
    inventario_cpu = input("Ingrese numero de inventario del cpu: ")

    system_info = get_system_info()
    monitor_info_list, monitor_count = get_monitor_info()



    for i in range(monitor_count):
        inventario_monitor = input(f"Ingrese numero de inventario del monitor {i + 1}: ")
        monitor_info_list[i].update(
            {'Departamento': departamento, 'Nombre de Usuario': nombre_usuario, 'Inventario': inventario_monitor,
             'Cargo': cargo})
        # Asegúrate de que todas las columnas necesarias estén en los datos nuevos
        for key in ['No.', 'Departamento', 'Nombre de Usuario', 'Correo Corporativo', 'Cargo', 'Tipo', 'Marca',
                    'Modelo', 'Serie', 'Inventario', 'Caracteristicas', 'Inf. en GB', 'MAC', 'IP']:
            if key not in monitor_info_list[i]:
                monitor_info_list[i][key] = ''

    # Intenta leer los datos existentes
    try:
        df_existing = read_existing_document()  # Asegúrate de que esta función retorna un DataFrame
        df_existing.reset_index(drop=True, inplace=True)
    except Exception as e:
        print("No existe un archivo de referencia, se creara uno")
        df_existing = pd.DataFrame(
            columns=['No.', 'Departamento', 'Nombre de Usuario', 'Correo Corporativo', 'Cargo', 'Tipo', 'Marca',
                     'Modelo',
                     'Serie', 'Inventario', 'Caracteristicas', 'Inf. en GB', 'MAC', 'IP'])

    # Actualiza la información del sistema con los datos ingresados por el usuario
    system_info.update(
        {'Departamento': departamento, 'Nombre de Usuario': nombre_usuario, 'Inventario': inventario_cpu,
         'Cargo': cargo})


    # Asegúrate de que todas las columnas necesarias estén en los datos nuevos
    for key in ['No.', 'Departamento', 'Nombre de Usuario', 'Correo Corporativo', 'Cargo', 'Tipo', 'Marca', 'Modelo',
                'Serie',
                'Inventario', 'Caracteristicas', 'Inf. en GB', 'MAC', 'IP']:
        if key not in system_info:
            system_info[key] = ''

    # Ahora crea un nuevo DataFrame con la información del sistema y del monitor

    monitor_info_list.insert(0, system_info)  # agrega system_info al inicio de monitor_info_list
    df_new = pd.DataFrame(monitor_info_list)
    df_existing = df_existing.reset_index(drop=True)
    df_existing.columns = df_existing.columns.str.upper()

    # Intenta reordenar las columnas ahora
    try:
        df_new = reorder_columns(df_new)
        df_new.columns = df_new.columns.str.upper()
        df_new = df_new.reset_index(drop=True)
        # print(df_new.columns)
    except Exception as e:
        print(f"Error al reordenar columnas: {e}")

    df_final = pd.concat([df_existing, df_new], ignore_index=True)

    # Auto incrementa y muestra el DataFrame final
    # try:
    #   df_final = auto_increment(df_final)
    #  print(df_final.columns)
    # except Exception as e:
    #   print(f"Error al auto incrementar: {e}")
    # Resetear el índice del DataFrame final

    df_final.reset_index(drop=True, inplace=True)

    # Renumerar la columna 'No.' basándote en el nuevo índice
    df_final['NO.'] = df_final.index + 1
    show_df(df_final)

    # Crea una ventana raíz y ocúltala
    root = Tk()
    root.withdraw()

    # Pregunta al usuario si desea guardar el documento
    save_doc = messagebox.askyesno("Guardar documento word", "¿Desea guardar en DOCX?")

    if save_doc:
        try:
            # Crea el documento final
            doc = Document()
            configure_margins(doc)
            add_title(doc)

            t = add_table_to_doc(df_final, doc)

            stylize_table(t)
            stylize_table_header(t)

            doc.save('DIRECCION DE TECNOLOGIAS.docx')
            messagebox.showinfo("Documento creado", "Documento creado con éxito")
        except (IOError, PermissionError) as e:
            # Muestra el mensaje de error si ocurre una excepción
            messagebox.showerror("Error", f"Ha ocurrido un error al crear el documento docx: {str(e)}")

    else:
        messagebox.showinfo("Creación de Documento Cancelada", "El usuario ha decidido no crear el archivo docx.")
    save_xlsx = messagebox.askyesno("Guardar documento excel", "¿Desea guardar en xlsx?")

    if save_xlsx:
        try:
            df_final = clean_illegal_chars(df_final)
            # Escribe el DataFrame en un archivo .xlsx
            save_to_excel(df_final)
            messagebox.showinfo("Documento creado", "Documento creado con éxito")
        except (IOError, PermissionError) as e:
            # Muestra el mensaje de error si ocurre una excepción
            messagebox.showerror("Error", f"Ha ocurrido un error al crear el documento xlsx: {str(e)}")
    else:
        messagebox.showinfo("Creación de Documento Cancelada", "El usuario ha decidido no crear el archivo xlsx.")

    root.destroy()


if __name__ == "__main__":
    main()
